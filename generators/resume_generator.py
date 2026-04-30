from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os

class ResumeGenerator:
    SUPPORTED_FORMATS = {'docx', 'pdf', 'txt'}

    def generate(self, content, output_format, file_path=None):
        if output_format not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {output_format}. Choose from {self.SUPPORTED_FORMATS}")

        if output_format == 'docx':
            return self._generate_docx(content, file_path)
        elif output_format == 'pdf':
            return self._generate_pdf(content, file_path)
        elif output_format == 'txt':
            return self._generate_txt(content, file_path)

    def _generate_docx(self, content, file_path=None):
        doc = Document()

        self._add_styled_paragraph(doc, 'OPTIMIZED RESUME', bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        lines = content.split('\n')
        in_heading = False

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.isupper() and len(stripped) < 60:
                p = doc.add_heading(stripped, level=1)
            elif stripped.startswith('-') or stripped.startswith('*'):
                doc.add_paragraph(stripped.lstrip('-* '), style='List Bullet')
            else:
                doc.add_paragraph(stripped)

        if file_path:
            doc.save(file_path)
            return file_path

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        return temp_file.name

    def _generate_pdf(self, content, file_path=None):
        try:
            from fpdf import FPDF
        except ImportError:
            raise ImportError("fpdf is required for PDF generation. Install with: pip install fpdf")

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "OPTIMIZED RESUME", ln=True, align="C")
        pdf.ln(5)

        pdf.set_font("Helvetica", "", 10)
        lines = content.split('\n')

        for line in lines:
            stripped = line.strip()
            if not stripped:
                pdf.ln(3)
                continue

            if stripped.isupper() and len(stripped) < 60:
                pdf.set_font("Helvetica", "B", 12)
                pdf.cell(0, 8, stripped, ln=True)
                pdf.ln(2)
            elif stripped.startswith('-') or stripped.startswith('*'):
                pdf.set_font("Helvetica", "", 10)
                pdf.cell(5)
                pdf.multi_cell(0, 5, f"  - {stripped.lstrip('-* ')}")
            else:
                pdf.set_font("Helvetica", "", 10)
                pdf.multi_cell(0, 5, stripped)

        if file_path:
            pdf.output(file_path)
            return file_path

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        pdf.output(temp_file.name)
        return temp_file.name

    def _generate_txt(self, content, file_path=None):
        if file_path:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return file_path

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode='w', encoding='utf-8')
        temp_file.write(content)
        temp_file.close()
        return temp_file.name

    def _add_styled_paragraph(self, doc, text, bold=False, size=12, alignment=None):
        p = doc.add_paragraph()
        if alignment:
            p.alignment = alignment
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p
